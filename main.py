# This is a Python script used to fill up a template docx file.
import sys

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


def generate_docx(template_file_name, person_info, paragraphs, font_name="FangSong"):
    (student, course, time) = person_info

    outfile = (student.replace('\n', '') + '-'
               + course.replace('\n', '') + '-'
               + (time + "").split(' ')[0] + '.docx')
    print("outfile name: " + outfile)

    docx = Document(template_file_name)

    # reset the default font in the theme
    style = docx.styles['Normal']
    font = style.font
    font.name = font_name
    font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    docx.save(outfile)

    docx = Document(outfile)

    table = docx.tables[0]

    table.cell(0, 1).text = "学生：" + student
    table.cell(0, 2).text = "科目：" + course
    table.cell(1, 0).text = "时间：" + time

    # remove the old content
    table.cell(2, 0).paragraphs.clear()
    for paragraph in paragraphs:
        table.cell(2, 0).add_paragraph(paragraph)

    docx.save(outfile)
    print("updated")


def parse_input(person_info, paragraphs):
    filename = args[1]
    with open(filename) as f:
        paragraph = ''

        while True:
            line = f.readline()
            if not line:
                paragraphs.append(paragraph)
                break
            # process(line)
            if len(person_info) < 3:
                person_info.append(line)
            else:
                if len(line) == 0:
                    if paragraph != '':
                        paragraphs.append(paragraph)
                    paragraph = ''
                else:
                    paragraph += line


if __name__ == '__main__':
    args = sys.argv
    print(args)

    if len(args) < 2:
        print("usage: python main.py input_file")
        exit(1)

    person_info_list = []
    paragraph_list = []
    parse_input(person_info_list, paragraph_list)

    generate_docx("demo.docx", person_info_list, paragraph_list)

    print("done")
